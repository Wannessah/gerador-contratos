# Anotar o passo a passo de como resolver esse problema
# quais tecnicas podem resolver essa demanda
# ler planilha openpyxl
# criar arquivo word: python-docx

# anotar os passos manuais que podem ser transformado em código

from openpyxl import load_workbook
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os

# Função para criar o contrato
def criar_contrato(nome_empresa, endereco, cidade, estado, cep, email):
    documento = Document()
    
    # Adicionar título centralizado
    titulo = documento.add_heading('Contrato de Prestação de Serviço', level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Adicionar parágrafo inicial com detalhes do fornecedor
    paragrafo_inicial = documento.add_paragraph()
    paragrafo_inicial.add_run(f'Este contrato de prestação de serviços é feito entre {nome_empresa}, com endereço em {endereco}, {cidade}, {estado}, CEP {cep}, doravante denominado FORNECEDOR, e a empresa CONTRATANTE.\n')

    # Adicionar corpo do contrato
    corpo = documento.add_paragraph()
    corpo.add_run('Pelo presente instrumento particular, as partes têm, entre si, justo e acordado o seguinte:\n\n')

    corpo.add_run('1. OBJETO DO CONTRATO\n').bold = True
    corpo.add_run('O FORNECEDOR compromete-se a fornecer à CONTRATANTE os serviços/material de acordo com as especificações acordadas, respeitando os padrões de qualidade e os prazos estipulados.\n\n')

    corpo.add_run('2. PRAZO\n').bold = True
    corpo.add_run('Este contrato tem prazo de vigência de 12 (doze) meses, iniciando-se na data de sua assinatura, podendo ser renovado conforme acordo entre as partes.\n\n')

    corpo.add_run('3. VALOR E FORMA DE PAGAMENTO\n').bold = True
    corpo.add_run('O valor dos serviços prestados será acordado conforme as demandas da CONTRATANTE e a capacidade de entrega do FORNECEDOR. Os pagamentos serão realizados mensalmente, mediante apresentação de nota fiscal.\n\n')

    corpo.add_run('4. CONFIDENCIALIDADE\n').bold = True
    corpo.add_run('Todas as informações trocadas entre as partes durante a vigência deste contrato serão tratadas como confidenciais.\n\n')

    # Adicionar assinaturas
    assinaturas = documento.add_paragraph()
    assinaturas.add_run('Para firmeza e como prova de assim haverem justo e contratado, as partes assinam o presente contrato em duas vias de igual teor e forma.\n\n')

    assinaturas.add_run(f'FORNECEDOR: {nome_empresa}\n')
    assinaturas.add_run(f'E-mail: {email}\n\n')
    assinaturas.add_run(f'CONTRATANTE: Prestador Sampa SA\n')
    assinaturas.add_run(f'E-mail: prestador_sampa_sa@gmail.com\n\n')
    assinaturas.add_run(f'São Paulo, {datetime.now().strftime("%d/%m/%Y")}\n')

    return documento

# Caminho para a planilha e pasta de saída
planilha_path = './fornecedores.xlsx'
pasta_contratos = './contratos'

# Verificar se a pasta de contratos existe, se não, criar
os.makedirs(pasta_contratos, exist_ok=True)

# Carregar a planilha Excel
planilha_fornecedores = load_workbook(planilha_path)
pagina_fornecedores = planilha_fornecedores['Sheet1']

# Iterar sobre as linhas da planilha
for linha in pagina_fornecedores.iter_rows(min_row=2, values_only=True):
    nome_empresa, endereco, cidade, estado, cep, telefone, email, setor = linha

    # Criar o contrato
    documento = criar_contrato(nome_empresa, endereco, cidade, estado, cep, email)
    
    # Salvar o contrato
    arquivo_path = os.path.join(pasta_contratos, f'contrato_{nome_empresa}.docx')
    documento.save(arquivo_path)
